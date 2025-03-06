import os
import csv
import json
import logging
import ssl

# Try to import optional dependencies with fallbacks
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    logging.warning("pdfplumber not installed. PDF parsing will be limited.")

try:
    import docx
except ImportError:
    docx = None
    logging.warning("python-docx not installed. DOCX parsing will be unavailable.")

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None
    logging.warning("pdf2image not installed. PDF image extraction will be unavailable.")
from tools.vector_search import VectorSearch
from tools.graph_search import GraphSearch
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer, SnowballStemmer
import string
import shutil
import logging
from graph.response_generation import ResponseGenerator
import redis

try:
    import pytesseract
except ImportError:
    pytesseract = None
    logging.warning("pytesseract not installed. OCR functionality will be unavailable.")

try:
    from langdetect import detect, LangDetectException
except ImportError:
    logging.warning("langdetect not installed. Language detection will be unavailable.")
    
    # Define fallback functions
    def detect(text):
        return "english"
        
    class LangDetectException(Exception):
        pass
from tools.memory_manager import MemoryManager
import aiofiles
import asyncio
from config.config import Config

# Add additional stopwords if available
try:
    # Fix SSL certificate issue
    ssl._create_default_https_context = ssl._create_unverified_context
    stopwords.words('german')
    stopwords.words('turkish')
except Exception as e:
    logging.warning(f"Could not load all NLTK stopwords: {e}")
    # Create fallback stopwords for essential languages
    FALLBACK_STOPWORDS = {
        'english': ['a', 'an', 'the', 'and', 'or', 'but', 'if', 'because', 'as', 'what', 
                  'which', 'this', 'that', 'these', 'those', 'then', 'just', 'so', 'than', 'such',
                  'when', 'while', 'who', 'with', 'you', 'your', 'yours', 'not', 'no', 'nor', 'none'],
        'german': ['der', 'die', 'das', 'ein', 'eine', 'und', 'oder', 'aber', 'wenn', 'weil', 'als'],
        'turkish': ['ve', 'ile', 'bu', 'şu', 'bir', 'için', 'gibi', 'de', 'da']
    }

class DocumentHandler:
    """Handles document uploads, indexing in vector search, and knowledge graph integration."""
    
    def __init__(self, storage_dir=Config.STORAGE_DIR, redis_url=Config.REDIS_URL, memory_manager=None, vector_search=None, graph_search=None, response_generator=None):
        self.storage_dir = storage_dir
        self.vector_search = vector_search or VectorSearch()
        self.graph_search = graph_search or GraphSearch()
        os.makedirs(self.storage_dir, exist_ok=True)
        self.version_dir = Config.VERSION_DIR
        os.makedirs(self.version_dir, exist_ok=True)
        
        # Create a test document if none exist
        try:
            self.documents_dir = self.storage_dir
            doc_files = [f for f in os.listdir(self.documents_dir) if os.path.isfile(os.path.join(self.documents_dir, f))]
            if not doc_files:
                test_path = os.path.join(self.documents_dir, "test_document.txt")
                with open(test_path, "w") as f:
                    f.write("This is a test document for the LangGraph system.\n\n")
                    f.write("You can use this document to test document retrieval functionality.\n")
                    f.write("The system should be able to find and retrieve this document when asked.\n")
                print(f"Created test document at {test_path}")
        except Exception as e:
            print(f"Error creating test document: {e}")
        
        self.response_generator = response_generator or ResponseGenerator()
        try:
            self.redis_client = redis.Redis.from_url(redis_url)
        except Exception as e:
            logging.warning(f"Could not connect to Redis: {e}")
            self.redis_client = None
        if memory_manager is None:
            from tools.memory_manager import initialize_database
            db_url = Config.SQLALCHEMY_DB_URL
            backend = initialize_database(db_url=db_url)
            self.memory_manager = MemoryManager(memory_backend=backend)
        else:
            self.memory_manager = memory_manager
        
        # Store pytesseract module if available
        if pytesseract is not None:
            self.pytesseract = pytesseract
            
        self._initialize_components()
    
    def _initialize_components(self):
        """Initialize all necessary components."""
        try:
            # Initialize stopwords if NLTK is working properly
            self.stop_words = {
                'english': set(stopwords.words('english')),
                'german': set(stopwords.words('german')),
                'turkish': set(stopwords.words('turkish'))
            }
        except Exception as e:
            logging.warning(f"Could not initialize NLTK stopwords: {e}. Using fallback stopwords.")
            # Use fallback stopwords
            if 'FALLBACK_STOPWORDS' in globals():
                self.stop_words = {lang: set(words) for lang, words in FALLBACK_STOPWORDS.items()}
            else:
                # Minimal fallback if global fallback wasn't defined
                self.stop_words = {
                    'english': set(['a', 'an', 'the', 'and', 'or', 'but', 'if']),
                    'german': set(['der', 'die', 'das', 'und', 'oder']),
                    'turkish': set(['ve', 'ile', 'bu'])
                }
                
        # Initialize stemmers
        try:
            self.porter_stemmer = PorterStemmer()
            self.snowball_stemmer = SnowballStemmer('english')
        except Exception as e:
            logging.warning(f"Could not initialize stemmers: {e}. Text processing will be limited.")
            # Create simple stemming functions as fallbacks
            self.porter_stemmer = self.snowball_stemmer = type('SimpleStemmer', (), {'stem': lambda self, word: word.lower()})()
            
        # Initialize other components if needed
        # ...additional initialization code...

    def preprocess_text(self, text, language='english'):
        # Use class-level stop words to handle fallbacks
        if language not in self.stop_words:
            language = 'english'  # Default to English if language not supported
            
        stop_words = self.stop_words[language]
        
        # Use class-level stemmers for consistency
        stemmer = self.porter_stemmer if language == 'english' else self.snowball_stemmer
        
        try:
            # Use word_tokenize if available
            words = word_tokenize(text)
        except Exception as e:
            # Basic fallback tokenization if NLTK fails
            logging.warning(f"NLTK word_tokenize failed: {e}. Using basic tokenization.")
            words = text.split()
            
        # Filter words
        filtered_words = []
        for word in words:
            word_lower = word.lower()
            if word_lower not in stop_words and word not in string.punctuation:
                try:
                    stemmed_word = stemmer.stem(word)
                    filtered_words.append(stemmed_word)
                except Exception:
                    # If stemming fails, just use the original word
                    filtered_words.append(word_lower)
                    
        return ' '.join(filtered_words)

    def detect_language(self, text):
        try:
            return detect(text)
        except LangDetectException:
            logging.error("Language detection failed.")
            return 'english'  # default to English if detection fails

    async def generate_summary(self, text, ratio=0.2):
        try:
            execution_result = "Summarize the following text in English language."
            retrieved_context = text
            return await self.response_generator.generate(execution_result, retrieved_context)
        except Exception as e:
            return f"An error occurred while generating summary: {str(e)}"

    async def save_document(self, file_path):
        try:
            file_name = os.path.basename(file_path)
            destination = os.path.join(self.storage_dir, file_name)
            version_path = self._save_version(file_path)
            text = await self._extract_text(file_path)
            if text is None:
                raise RuntimeError(f"Error processing file: {file_path}")
            async with aiofiles.open(destination.replace(os.path.splitext(file_path)[1], ".txt"), "w", encoding="utf-8") as f:
                await f.write(text)
            language = self.detect_language(text)
            preprocessed_text = self.preprocess_text(text, language)
            summary = await self.generate_summary(preprocessed_text)
            summary_path = destination.replace(".txt", "_summary.txt")
            async with aiofiles.open(summary_path, "w", encoding="utf-8") as f:
                await f.write(summary)
            text_chunks = self.chunk_text(preprocessed_text)
            metadata = await self._extract_metadata(file_path)
            yield 0.33
            for chunk in text_chunks:
                self.vector_search.add_document(file_name, chunk, metadata)
            yield 0.66
            self.graph_search.add_relationship("Documents", file_name, "contains", metadata)
            yield 1.0
            self._identify_and_store_relationships(file_name, preprocessed_text)
            
            # Update knowledge graph with new document information
            self.graph_search.add_knowledge("Documents", "contains", file_name, metadata)
            
            yield f"Document '{file_name}' processed and stored."
        except FileNotFoundError:
            logging.error(f"Error: The file '{file_path}' was not found.")
            yield f"Error: The file '{file_path}' was not found."
        except Exception as e:
            logging.error(f"An unexpected error occurred: {str(e)}")
            yield f"An unexpected error occurred: {str(e)}"
    
    def _save_version(self, file_path):
        file_name = os.path.basename(file_path)
        version_number = self._get_next_version_number(file_name)
        version_file_name = f"{file_name}_v{version_number}"
        version_path = os.path.join(self.version_dir, version_file_name)
        shutil.copy(file_path, version_path)
        return version_path
    
    def _get_next_version_number(self, file_name):
        existing_versions = [f for f in os.listdir(self.version_dir) if f.startswith(file_name)]
        return len(existing_versions) + 1
    
    def list_versions(self, file_name):
        return [f for f in os.listdir(self.version_dir) if f.startswith(file_name)]
    
    def retrieve_version(self, file_name, version_number):
        version_file_name = f"{file_name}_v{version_number}"
        version_path = os.path.join(self.version_dir, version_file_name)
        try:
            if os.path.exists(version_path):
                with open(version_path, "r", encoding="utf-8") as f:
                    return f.read()
            logging.error(f"Document version '{version_file_name}' not found.")
            return "Document version not found."
        except FileNotFoundError:
            logging.error(f"Error: The file '{version_file_name}' was not found.")
            return f"Error: The file '{version_file_name}' was not found."
        except Exception as e:
            logging.error(f"An unexpected error occurred: {str(e)}")
            return f"An unexpected error occurred: {str(e)}"
    
    async def _extract_text(self, file_path):
        try:
            if file_path.endswith(".pdf"):
                return await self._extract_text_from_pdf(file_path)
            elif file_path.endswith(".docx"):
                return await self._extract_text_from_docx(file_path)
            elif file_path.endswith(".csv"):
                return await self._extract_text_from_csv(file_path)
            elif file_path.endswith(".json"):
                return await self._extract_text_from_json(file_path)
            else:
                async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                    return await f.read()
        except Exception as e:
            logging.error(f"Error processing file '{file_path}': {str(e)}")
            return None
    
    async def _extract_text_from_pdf(self, pdf_path):
        text = ""
        try:
            if pdfplumber is None:
                logging.error("pdfplumber is not installed. Cannot extract text from PDF.")
                return "PDF extraction is not available. Please install pdfplumber."
                
            with pdfplumber.open(pdf_path) as pdf:
                for page_number, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    elif convert_from_path is not None and hasattr(self, 'pytesseract'):
                        # Only try image extraction if pdf2image and pytesseract are available
                        image = convert_from_path(pdf_path, first_page=page_number+1, last_page=page_number+1)[0]
                        text += self.pytesseract.image_to_string(image) + "\n"
                    else:
                        text += "[Image content not extractable without pdf2image/pytesseract]\n"
        except FileNotFoundError:
            logging.error(f"Error: The file '{pdf_path}' was not found.")
            raise FileNotFoundError(f"Error: The file '{pdf_path}' was not found.")
        except Exception as e:
            logging.error(f"An unexpected error occurred while extracting text from PDF: {str(e)}")
            raise Exception(f"An unexpected error occurred while extracting text from PDF: {str(e)}")
        return text

    async def _extract_text_from_docx(self, docx_path):
        text = ""
        try:
            if docx is None:
                logging.error("python-docx is not installed. Cannot extract text from DOCX.")
                return "DOCX extraction is not available. Please install python-docx."
                
            doc = docx.Document(docx_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except FileNotFoundError:
            logging.error(f"Error: The file '{docx_path}' was not found.")
            raise FileNotFoundError(f"Error: The file '{docx_path}' was not found.")
        except Exception as e:
            logging.error(f"An unexpected error occurred while extracting text from DOCX: {str(e)}")
            raise Exception(f"An unexpected error occurred while extracting text from DOCX: {str(e)}")
        return text

    async def _extract_text_from_csv(self, csv_path):
        text = ""
        try:
            async with aiofiles.open(csv_path, "r", encoding="utf-8") as csvfile:
                reader = csv.reader(await csvfile.read())
                for row in reader:
                    text += ", ".join(row) + "\n"
        except FileNotFoundError:
            logging.error(f"Error: The file '{csv_path}' was not found.")
            raise FileNotFoundError(f"Error: The file '{csv_path}' was not found.")
        except Exception as e:
            logging.error(f"An unexpected error occurred while extracting text from CSV: {str(e)}")
            raise Exception(f"An unexpected error occurred while extracting text from CSV: {str(e)}")
        return text

    async def _extract_text_from_json(self, json_path):
        text = ""
        try:
            async with aiofiles.open(json_path, "r", encoding="utf-8") as jsonfile:
                data = json.load(await jsonfile.read())
                text = json.dumps(data, indent=4)
        except FileNotFoundError:
            logging.error(f"Error: The file '{json_path}' was not found.")
            raise FileNotFoundError(f"Error: The file '{json_path}' was not found.")
        except Exception as e:
            logging.error(f"An unexpected error occurred while extracting text from JSON: {str(e)}")
            raise Exception(f"An unexpected error occurred while extracting text from JSON: {str(e)}")
        return text
    
    async def _extract_metadata(self, file_path):
        metadata = {
            "author": "Unknown",
            "title": os.path.basename(file_path),
            "keywords": []
        }
        if file_path.endswith(".pdf"):
            metadata.update(await self._extract_metadata_from_pdf(file_path))
        elif file_path.endswith(".docx"):
            metadata.update(await self._extract_metadata_from_docx(file_path))
        return metadata
    
    async def _extract_metadata_from_pdf(self, pdf_path):
        metadata = {}
        try:
            if pdfplumber is None:
                logging.error("pdfplumber is not installed. Cannot extract metadata from PDF.")
                return {
                    "author": "Unknown",
                    "title": os.path.basename(pdf_path),
                    "keywords": []
                }
                
            with pdfplumber.open(pdf_path) as pdf:
                doc_info = pdf.metadata
                metadata["author"] = doc_info.get("Author", "Unknown")
                metadata["title"] = doc_info.get("Title", os.path.basename(pdf_path))
                metadata["keywords"] = doc_info.get("Keywords", "").split(", ")
        except Exception as e:
            logging.error(f"An error occurred while extracting metadata from PDF: {str(e)}")
        return metadata
    
    async def _extract_metadata_from_docx(self, docx_path):
        metadata = {}
        try:
            if docx is None:
                logging.error("python-docx is not installed. Cannot extract metadata from DOCX.")
                return {
                    "author": "Unknown",
                    "title": os.path.basename(docx_path),
                    "keywords": []
                }
                
            doc = docx.Document(docx_path)
            core_props = doc.core_properties
            metadata["author"] = core_props.author or "Unknown"
            metadata["title"] = core_props.title or os.path.basename(docx_path)
            metadata["keywords"] = core_props.keywords.split(", ") if core_props.keywords else []
        except Exception as e:
            logging.error(f"An error occurred while extracting metadata from DOCX: {str(e)}")
        return metadata
    
    def list_documents(self):
        try:
            return [f for f in os.listdir(self.storage_dir) if f.endswith(".txt") or f.endswith(".pdf")]
        except Exception as e:
            logging.error(f"An error occurred while listing documents: {str(e)}")
            return f"An error occurred while listing documents: {str(e)}"
    
    def retrieve_document(self, filename):
        """Retrieves a document from cache or file system."""
        try:
            cache_key = f"document:{filename}"
            content = self.memory_manager.get_cache(cache_key)
            if content is None:
                file_path = os.path.join(self.storage_dir, filename)
                if os.path.exists(file_path):
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()
                    self.memory_manager.set_cache(cache_key, content, ttl=Config.CACHE_TTL)  # Cache for 1 day
                else:
                    logging.error(f"Document '{filename}' not found.")
                    return "Document not found."
            return content
        except FileNotFoundError:
            logging.error(f"Error: The file '{filename}' was not found.")
            return f"Error: The file '{filename}' was not found."
        except Exception as e:
            logging.error(f"An unexpected error occurred: {str(e)}")
            return f"An unexpected error occurred: {str(e)}"
    
    def is_document_query(self, query):
        keywords = ["document", "file", "pdf", "txt"]
        return any(keyword in query.lower() for keyword in keywords)
    
    def _identify_and_store_relationships(self, file_name, preprocessed_text):
        similar_docs = self.vector_search.find_similar_documents(preprocessed_text)
        for similar_doc in similar_docs:
            self.graph_search.add_relationship("Documents", file_name, "related_to", similar_doc)
            
    async def handle_document(self, state):
        """Handler for document operations in the workflow graph.
        
        This method is called by the workflow graph to process document-related
        operations based on the current state.
        
        Args:
            state: The current workflow state containing user query and context
            
        Returns:
            Updated state with document handling results
        """
        try:
            # Extract document handling intent from query
            user_query = state.get("user_query", "")
            
            # Check if this is a document query
            if not self.is_document_query(user_query):
                # Not a document query, skip processing
                return state
                
            # Parse the query to identify the operation and document name
            operation = "retrieve"  # Default operation
            
            if "upload" in user_query.lower() or "save" in user_query.lower():
                operation = "save"
            elif "list" in user_query.lower():
                operation = "list"
            elif "search" in user_query.lower():
                operation = "search"
                
            # Process based on operation type
            if operation == "list":
                # List available documents
                documents = self.list_documents()
                state["execution_result"] = {"document_list": documents}
                
            elif operation == "search":
                # Extract search terms from query
                search_terms = user_query.lower().split()
                search_terms = [term for term in search_terms if term not in ["search", "document", "for", "about", "find"]]
                search_query = " ".join(search_terms)
                
                # Use vector search to find relevant documents
                results = self.vector_search.search_similar(search_query)
                state["execution_result"] = {"document_search_results": results}
                
            elif operation == "save":
                # Note: In a real implementation, you'd need to handle the file path
                # For now, just acknowledge the intent
                state["execution_result"] = {"document_save_result": "Document save request acknowledged"}
                
            else:  # Default to retrieve
                # Extract document name
                file_name = self._extract_document_name(user_query)
                
                if file_name:
                    # Retrieve the document
                    content = self.retrieve_document(file_name)
                    state["execution_result"] = {"document_content": content}
                else:
                    state["execution_result"] = {"document_error": "Could not determine which document to retrieve"}
                    
            return state
            
        except Exception as e:
            logging.error(f"Error in handle_document: {e}")
            state["error"] = f"Document handling error: {str(e)}"
            return state
